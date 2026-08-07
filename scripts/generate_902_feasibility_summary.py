#!/usr/bin/env python3
"""Generate 902 PDF Extension Feasibility Study summary Word document."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def build_document():
    doc = Document()
    title = doc.add_heading(
        "凯莱英 UK Sandwich 902 厂房扩建项目\nFeasibility Study 详细汇总",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        f"编制说明：结合 Scitech-EKIUM 交付包（Expension_Feasible_Study_2026-5-27）\n"
        f"及 2026-05-27 可行性研究讲解会议转写整理\n"
        f"整理日期：{date.today().isoformat()}"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(80, 80, 80)

    add_heading(doc, "一、项目背景与定位", 1)
    add_bullets(
        doc,
        [
            "项目地点：Discovery Park（原 Pfizer 园区），Kent Sandwich，现有 B902 PDF（Pilot Development Facility）厂房东侧扩建。",
            "设计方：Scitech-EKIUM（合同方 Asymchem），阶段为 RIBA Stage 1 Feasibility Study（Issue A1，2026-05-22）。",
            "扩建目标：在现有厂房旁新建约 600 m² 四层主体 + 顶层设备夹层（Plant Mezzanine），形成约 3,099 m² GIFA，用于安装 10 台反应釜（R25–R34）、1 台 2500 L 加氢釜（R35）、3 套 Filter Dryer、物料处理/连续制造区及配套 HVAC、公用工程与自控。",
            "与 902 既有改造（HPLC/冻干）为不同子项目：本汇总仅针对 902 东侧新楼扩建；HPLC/冻干为既有厂房内改造，由设备管理部主导。",
        ],
    )

    add_heading(doc, "二、2026-05-27 会议要点解读（Scitech 讲解）", 1)

    add_heading(doc, "2.1 会议性质与共识", 2)
    add_bullets(
        doc,
        [
            "无单独 PPT，逐章走读 Feasibility Report；允许录音形成纪要。",
            "明确：Feasibility 深度有限，用于判断“是否值得进入下一阶段”，非施工图级设计。",
            "执行摘要结论（与报告一致）：在取得当地规划许可前提下，902 东侧扩建方案可行（feasible）。",
            "部分 B902 运行数据仍缺失，报告含假设；工艺/机械信息需在 RIBA Stage 2 通过 workshop 继续澄清。",
        ],
    )

    add_heading(doc, "2.2 建筑与方案选项（会议重点）", 2)
    add_bullets(
        doc,
        [
            "推荐基线方案（Option 1，评分最高）：拆除/迁址现有 Hydrogenation 小厂房；新楼与 902 地面、一、二层楼面齐平贯通；三层为加氢区并整体抬高约 2 m；三层上方增设设备夹层（不覆盖加氢房间）。",
            "若保留原加氢厂房：需整体前移新楼或缩小建筑面积，中央走廊对称性变差，造价与施工安全变差——矩阵评分均低于 Option 1。",
            "曾评估“整体抬高楼板 + 剪刀升降机”方案（Option 4）：层间高差过大（约 4–5 m 台阶），运维不可接受，已否决。",
            "首层 3 台 Hastelloy 反应釜净高风险：可利用场地已有约 1.5 m 防洪抬升，在首层局部降低楼板 1–1.5 m 增净高；若仍不足，Stage 2 再评估是否上移至楼上（工艺上首层更合理）。",
            "货梯位置调整：从扩建楼内移出，布置在旧楼与新楼之间、正对主楼梯——释放 I/O 与机房空间（与报告 5.2 节一致）。",
            "三层需新增楼梯+平台升降机，连接旧 PDF 三层与新扩建功能区。",
            "加氢区设 blast panel；新楼末端增设第二消防疏散楼梯。",
        ],
    )

    add_heading(doc, "2.3 场址、防洪与施工", 2)
    add_bullets(
        doc,
        [
            "River Stour 防洪：新楼地面约比室外地坪高 1.5 m，与 902、530 一致。",
            "基础：建议延续桩基础（与 902/530 类似），需补充钻孔；土壤条件一般。",
            "结构：延续钢结构体系，与现有楼一致，避免混凝土/钢结构混用导致不均匀变形。",
            "与 902 连接：东侧原有 knock-out panel 需勘测是否满足新楼标高；可能需增楼梯/电梯消化标高差。",
            "可施工性（Constructability）：场址可建；需关注与运行厂房的 interface、吊装通道、外部 LN₂/H₂ 罐区道路。",
            "未爆弹药（UXO）：二战航线下方，建议金属探测/专项调查——会议亦提及此条。",
        ],
    )

    add_heading(doc, "2.4 规划、BREEAM 与法规", 2)
    add_bullets(
        doc,
        [
            "规划：扩建需正式 Planning Consent；园区为 Enterprise Zone，流程相对简化，但仍须 pre-application 与当局沟通——Scitech 正在编制 Pre-Planning 文件。",
            "BREEAM：目标 Very Good；预评估约 62.93%（及格线 55%），尚有约 20 分可争取但可能增加造价（如屋面 PV 等）；REAP Stage 2 宜尽早与 BREEAM Assessor 对接。",
            "建筑规范：已预留防火疏散、现有排烟烟囱加高或改路由至新楼屋面等议题，Stage 2 由 Fire Consultant 统一修订消防策略。",
            "福利设施：人员 45→60，建议在旧楼将闲置 Disabled WC/Shower 改造为 2 间 Super-loo，并评估更衣/淋浴/休息区（报告第 6 章）。",
        ],
    )

    add_heading(doc, "2.5 公用工程（会议与报告一致要点）", 2)
    add_bullets(
        doc,
        [
            "规模跃迁：现有最大釜 2500 L → 扩建含 5000/8000 L 等；设备总容积 scale-up 约 3.54×。",
            "冷/热 HTF：稳态负荷下现有 780 kW 氨冷机约 69% 利用率；crash cooling 待 Stage 2 用放热数据校核。",
            "氢气：扩建 2500 L 加氢釜按约 10× 估算，建议外置 MCP 储氢，叉车可达；可与现有气瓶库统筹。",
            "液氮：建议新建约 35 m³ 罐（Option 2），6.5×10.5 m 基础；可用 artic 槽车，年节省配送费约 £50k。",
            "压缩空气：现有 375 Nm³/h 包机在可行性阶段可满足合并需求（约 80% 利用率）。",
            "PW：扩建需独立 PW 系统（发生、16 m³ 罐、25 m³/h 泵等），不宜仅换泵。",
            "溶剂/废液：现有储罐容积够，主要增加配送/转运频次；废液罐 ST01–03 需提高转移频率。",
            "电力：会议提及现有变压器容量经初步核算可覆盖扩建增量（需与报告 10 章负荷表一并内部复核）。",
            "自控：报告建议 DeltaV 扩展集成（11 章）。",
        ],
    )

    add_heading(doc, "2.6 费用与进度（会议口径）", 2)
    add_bullets(
        doc,
        [
            "可行性 OOM 总造价（含风险与预备费）：约 £78.1M（2026-05-22 基准，非最终 Capex）。",
            "其中：工程直接费小计约 £42.0M；工艺设备包约 £23.2M；基础建造成本约 £47.6M；设计/勘测/BREEAM/规划等约 £4.1M；风险登记册量化 £0.63M；设计发展/施工/业主预备费合计约 £25.8M。",
            "说明：不含业主采购工艺设备（Client purchased equipment excluded），但含吊装就位等 allowance。",
            "下一阶段：RIBA Stage 2 Concept Design，2026-07-01 起（见 Next Stage Programme）；每阶段预留约 1 个月 Client Review。",
            "风险登记册 #13：期望 2028 Q3 完工目标偏紧；正式计划显示 RIBA 5 施工 2028-08 起至 2029-11（需与 UK 团队确认里程碑定义）。",
            "长周期设备（Hastelloy 釜、Filter Dryer 等）宜在 Stage 4 前启动 18 个月采购；AHU 等约 4 个月。",
        ],
    )

    add_heading(doc, "三、Feasibility 报告技术摘要", 1)

    add_heading(doc, "3.1 建筑面积与功能分区", 2)
    add_bullets(
        doc,
        [
            "GIFA 合计 3,099 m²：Ground 597；First 642；Second 637；Third + Mezz 1,224。",
            "功能占比（约）：Primary/Operational 1,301 m²；Plant/Technical 1,168 m²；Circulation 402 m² 等。",
            "首层（示例）：PW 系统、开关间、Milling & Packaging、Hastelloy ADF、Continuous 区等。",
            "二层：多台玻璃衬里/Hastelloy 反应釜（R25–R34）及冷凝、夹套循环等。",
            "三层：加氢 suite（R35）及 blast 设计；设备夹层布置 AHU/冷机等。",
        ],
    )

    add_heading(doc, "3.2 主要工艺设备（Equipment List A1 摘要）", 2)
    add_bullets(
        doc,
        [
            "反应釜 R25–R31：多为 5000 L 玻璃衬里，位于 2nd floor；R26/R28/R34 等为 8000 L。",
            "R32/R33：Hastelloy，5000/8000 L；R33 具备氢气排放能力（IIC T4）选项。",
            "R35：2500 L working（3400 L gross）加氢釜，3rd floor，设计压力至 2.4 MPa。",
            "Filter Dryer FD01–03、Distillate Receivers RC01–05、Milling（Corn mill）等见附录设备表。",
        ],
    )

    add_heading(doc, "3.3 HVAC 与机电（概要）", 2)
    add_bullets(
        doc,
        [
            "6 台通用 AHU + 1 台 Filter Dryer/Milling/Packing 专用 AHU；LEV、溶剂转移排风、热回收（峰值可节能约 75% 冷暖负荷）。",
            "3 台约 100 kW 冷机、6 台 AHU、Sprinkler 延伸、Steam→LTHW 换热撬、BMS 接入园区系统。",
            "电气：新主配电、300 kVA 柴油发电机、UPS、Lift、防雷、CCTV/门禁/Interlock 等。",
        ],
    )

    add_heading(doc, "3.4 主要风险（Risk Register 摘录）", 2)
    add_bullets(
        doc,
        [
            "#1 若不能拆除原 Hydrogenation 厂房 → 需重新做可行性研究（高）。",
            "#13 2028 Q3 完工目标紧 → 建议 D&B、24/7 施工等压缩策略（高，费用未量化）。",
            "#14 石棉 → Stage 3 前 R&D 调查（中，约 £80k 量级）。",
            "#16 设备涨价（高）；#17 不可抗力（高）；#18 采购策略不清（高）→ 需 Asymchem 明确职责与时间。",
            "#19 概念后设备范围变更（中）；#20 地表水排水与建筑冲突（中）。",
            "Knock-out panel 标高、HV 电缆走向、地下障碍物等须在 Stage 2–3 调查关闭。",
        ],
    )

    add_heading(doc, "四、建议的 RIBA Stage 2 行动清单（供国内统筹）", 1)
    add_bullets(
        doc,
        [
            "【决策】确认 Hydrogenation 旧厂房拆除/迁址方案及时间节点（影响总图与造价）。",
            "【决策】确认首层 Hastelloy 净高：降板 vs 设备上楼；同步工艺物流模拟。",
            "【信息】补齐 B902 运行数据、放热/crash cooling、阀门仪表清单、采购分工（Client vs Scitech）。",
            "【对接】启动 Pre-Planning 当局预审；BREEAM REAP Stage 2 工作坊。",
            "【对接】Fire Consultant 招标；UXO/地质/石棉等勘测招标。",
            "【商务】内部评审 £78.1M OOM 与资金节奏；长周期设备清单与 18 个月 lead time 对齐。",
            "【组织】设备管理部（郭宏杰/张暴）工艺需求冻结节点；TJ4（杨道兴/魏鹏）可参与操作经验与供应商接口（本扩建以反应/加氢/过滤为主，与 HPLC/冻干供应商并行管理）。",
        ],
    )

    add_heading(doc, "五、交付文件清单（GitHub 包内）", 1)
    add_bullets(
        doc,
        [
            "300291-RE-0001 Feasibility Report（72 页）",
            "300291-CM-0001 Feasibility Cost Plan",
            "300291-PM-PR-0002 Next Stage Programme",
            "300291-PM-RA-0001 Project Risk Register",
            "Option Appraisal Matrix；Equipment List；Utility Calculations",
            "各层平面 GA、剖面、AHU Zoning、BREEAM Pre-Assessment 等",
        ],
    )

    add_heading(doc, "六、关键数字速查", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    data = [
        ("可行性结论", "可行（以规划批准为前提）"),
        ("GIFA", "约 3,099 m²"),
        ("OOM 总造价", "约 £78.1M（2026-05-22，含预备费）"),
        ("工程费小计", "约 £41.97M"),
        ("工艺设备包", "约 £23.25M"),
        ("BREEAM 预评估", "约 62.93%（目标 Very Good）"),
        ("人员增加", "+15（45→60）"),
        ("Stage 2 启动", "2026-07-01（计划）"),
    ]
    for k, v in data:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v

    doc.add_page_break()
    add_heading(doc, "附录：会议与报告差异需跟进项", 1)
    add_bullets(
        doc,
        [
            "会议口语中“10 reactors”含加氢釜口径，设备表为 R25–R34（10 台）+ R35（加氢）——对外材料建议统一标签。",
            "报告写 R36 氢气 scale-up 处与 R35 标签需工艺专业核对（Utility 章节笔误可能）。",
            "完工日期：风险登记册写 2028 Q3 紧，总控计划施工至 2029 Q4——需 UK 项目方澄清“竣工/投料/验证”定义。",
        ],
    )

    return doc


def main():
    out_dir = Path("/workspace/会议")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "902扩建FeasibilityStudy详细汇总_2026-05-28.docx"
    doc = build_document()
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
