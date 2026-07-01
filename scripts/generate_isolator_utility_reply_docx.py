#!/usr/bin/env python3
"""Generate CN/EN Word summary for Clare — isolator utility requirements."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path("/workspace/会议/英国高活实验室隔离器公用工程信息_答复Clare_中英对照.docx")

NAVY = RGBColor(0x0F, 0x2B, 0x46)
GREY = RGBColor(0x5A, 0x6A, 0x7A)


def set_doc_font(doc, name="Microsoft YaHei"):
    style = doc.styles["Normal"]
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(10.5)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    r.italic = True


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    set_doc_font(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("英国高活实验室隔离器公用工程信息\nIsolator Utility Requirements — CN/EN Summary")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("For Woodley Cole HIPO Lab Feasibility Study  |  July 2026")
    r2.font.size = Pt(10)
    r2.font.color.rgb = GREY
    doc.add_paragraph()

    add_heading(doc, "问题 1 / Question 1", 2)
    add_heading(doc, "实验室内需为隔离器配套提供的公用工程", 3)
    add_note(doc, "Utility services to be provided in the lab (to operate the isolators)")

    q1_rows = [
        (
            "自来水 / Tap water",
            "需要；用于隔离器箱体内部清洗",
            "Required for internal chamber cleaning",
        ),
        (
            "电力 / Electrical power",
            "需要；实验室侧通常 220 V / 50 Hz 单路主电源。隔离器本体一般 ≤1.5 kW；总容量需叠加箱内设备（天平、小型反应/搅拌设备、油浴等）。集团内实验室参考约 1.2–1.3 kW；复杂场景可参考 ≤5.5 kW。待操作清单确认后最终定稿并留余量",
            "Typically 220 V / 50 Hz single supply. Isolator unit ≤1.5 kW; total must include in-isolator equipment. Internal lab reference ~1.2–1.3 kW; up to ~5.5 kW for complex cases. Final rating TBC with equipment list + margin",
        ),
        (
            "氮气 / Nitrogen",
            "需要；用于吹扫干燥、惰化及箱内压力控制",
            "Required for purge drying, inerting and pressure control",
        ),
        (
            "压缩空气 / Compressed air",
            "实验室操作侧通常不需要；隔离器可预留接口。氮气与压缩空气一般二选一，高活实验室多用氮气",
            "Generally not required for routine lab ops; optional connection on isolator. Typically nitrogen OR compressed air — nitrogen preferred for HP",
        ),
        (
            "真空 / Vacuum",
            "需要；干燥等操作使用。优先接入集中真空系统；亦可使用小型真空泵。隔离器仅预留接口，不带真空过滤器",
            "Required for drying. Central lab vacuum preferred; small dedicated pump acceptable. Interface only — no vacuum filter on isolator",
        ),
        (
            "排风 / 暖通 / Exhaust & HVAC",
            "需要；排风量与箱内体积及目标负压（−30~−50 Pa 或 −100~−200 Pa）有关。参考约 50 m³/h/排风口，详细设计阶段确认",
            "Required from building services. Flow depends on volume and target ΔP. Reference ~50 m³/h per exhaust point — TBC at detailed design",
        ),
        (
            "冷/热介质 / Heating & cooling utilities",
            "不接入隔离器；通过带入设备（如油浴）实现，仅用电",
            "Not connected to isolator — via equipment brought into chamber (e.g. oil bath), electrically powered",
        ),
    ]
    add_table(
        doc,
        ["公用工程\nUtility", "中文要点\nSummary (ZH)", "English summary"],
        q1_rows,
        col_widths=[3.2, 6.0, 6.0],
    )

    ctx = doc.add_paragraph()
    ctx.add_run("操作背景 / Typical operations: ").bold = True
    ctx.add_run(
        "高活固体物料分装、吹扫干燥、高活物料投料及部分反应操作；不涉及隔离器自带冷热介质接口。\n"
        "HP solid dispensing, purge drying, material addition and some reaction steps; no heating/cooling utility connections to isolator."
    )

    doc.add_paragraph()
    add_heading(doc, "问题 2 / Question 2", 2)
    add_heading(doc, "隔离器内部已配置 / 可提供的服务与接口", 3)
    add_note(doc, "Services available within the isolators")

    q2_rows = [
        ("称量台 / Weighing station", "配置电子天平称量台", "Balance platform"),
        ("电源插座 / Power sockets", "箱内插座；通过隔离器 PLC 控制盘统一供电", "Internal sockets via isolator PLC control panel"),
        ("监测仪表 / Sensors", "温度、湿度、压差传感器", "Temperature, humidity and differential pressure sensors"),
        ("清洗接口 / Cleaning", "水气枪（箱体内部清洗）", "Water/gas gun for internal chamber cleaning"),
        ("定制配置 / Custom fittings", "可按项目需求增配（如挂铲钩等）", "e.g. tool hooks — per project requirements"),
        ("真空接口 / Vacuum", "预留真空接口（不含过滤器）", "Connection point only — no filter included"),
        ("气体接口 / Gas", "预留氮气/压缩空气接口（按方案选配）", "Nitrogen and/or compressed air — per design option"),
        (
            "网络 / Network",
            "常规项目箱内不预埋网线；预留外部穿线接口。分析应用如有电子化需求，可能需网络接口——待确认",
            "No hardwired network inside (standard builds); external cable pass-through. Network port for analytical use — TBC",
        ),
    ]
    add_table(
        doc,
        ["项目\nItem", "中文要点\nSummary (ZH)", "English summary"],
        q2_rows,
        col_widths=[3.2, 6.0, 6.0],
    )

    doc.add_paragraph()
    add_heading(doc, "备注 / Notes", 2)
    notes = [
        "英国三明治站点项目尚处可行性早期；以上基于凯莱英集团国内高活实验楼及实验室隔离器参考实践，供当前可行性测算使用。",
        "最终功率、排风量及内部配置待英国侧操作需求及箱内设备清单确认后定稿。",
        "研发与分析实验室隔离器配置总体相近。",
        "UK Sandwich project at early feasibility stage. Based on Asymchem group reference practice (HP lab isolators in China). Final loads and internal configuration TBC once UK operating requirements and equipment list are defined. R&D and analytical isolator configurations are broadly similar.",
    ]
    for n in notes:
        p = doc.add_paragraph(n, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(10)

    doc.add_paragraph()
    add_note(doc, "Source: Internal meeting 2026-07-01 — 英国高活实验室隔离器公用工程信息讨论")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
